"""
File processor for handling different document formats.
Supports PDF, DOCX, CSV, and TXT files.
"""

import os
import tempfile
from pathlib import Path
from typing import List, Optional
from langchain_community.document_loaders import (
    PyPDFLoader,
    CSVLoader,
    TextLoader,
)
from langchain_core.documents import Document
from app.utils.logger_setup import log


def process_text_file(file_path: str) -> List[Document]:
    """Process plain text file."""
    try:
        loader = TextLoader(file_path, encoding='utf-8')
        docs = loader.load()
        log.info(f"Processed text file: {file_path}, documents: {len(docs)}")
        return docs
    except Exception as e:
        log.error(f"Error processing text file {file_path}: {e}")
        raise


def process_pdf_file(file_path: str) -> List[Document]:
    """Process PDF file."""
    try:
        loader = PyPDFLoader(file_path)
        docs = loader.load()
        log.info(f"Processed PDF file: {file_path}, pages: {len(docs)}")
        return docs
    except Exception as e:
        log.error(f"Error processing PDF file {file_path}: {e}")
        raise


def process_csv_file(file_path: str) -> List[Document]:
    """Process CSV file."""
    try:
        loader = CSVLoader(file_path)
        docs = loader.load()
        log.info(f"Processed CSV file: {file_path}, rows: {len(docs)}")
        return docs
    except Exception as e:
        log.error(f"Error processing CSV file {file_path}: {e}")
        raise


def process_docx_file(file_path: str) -> List[Document]:
    """Process DOCX file using python-docx."""
    try:
        from docx import Document as DocxDocument
        
        docx_doc = DocxDocument(file_path)
        
        # Extract all paragraphs
        full_text = []
        for para in docx_doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in docx_doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))
        
        # Create document
        content = "\n\n".join(full_text)
        docs = [Document(
            page_content=content,
            metadata={"source": file_path, "file_type": "docx"}
        )]
        
        log.info(f"Processed DOCX file: {file_path}, content length: {len(content)}")
        return docs
    except Exception as e:
        log.error(f"Error processing DOCX file {file_path}: {e}")
        raise


def process_uploaded_file(file_path: str, file_name: str) -> List[Document]:
    """
    Process an uploaded file based on its extension.
    
    Args:
        file_path: Full path to the uploaded file
        file_name: Original filename with extension
        
    Returns:
        List of Document objects
        
    Raises:
        ValueError: If file type is not supported
    """
    file_ext = Path(file_name).suffix.lower()
    
    if file_ext == '.pdf':
        return process_pdf_file(file_path)
    elif file_ext == '.docx':
        return process_docx_file(file_path)
    elif file_ext == '.csv':
        return process_csv_file(file_path)
    elif file_ext in ['.txt', '.text']:
        return process_text_file(file_path)
    else:
        log.error(f"Unsupported file type: {file_ext}")
        raise ValueError(f"Unsupported file type: {file_ext}. Supported types: pdf, docx, csv, txt")


def cleanup_file(file_path: str) -> bool:
    """
    Delete a file after processing.
    
    Args:
        file_path: Path to file to delete
        
    Returns:
        True if successful, False otherwise
    """
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            log.info(f"Cleaned up file: {file_path}")
            return True
        return False
    except Exception as e:
        log.error(f"Error cleaning up file {file_path}: {e}")
        return False


def create_temp_upload_dir() -> str:
    """Create a temporary directory for file uploads."""
    temp_dir = tempfile.mkdtemp(prefix="rag_uploads_")
    log.info(f"Created temp upload directory: {temp_dir}")
    return temp_dir
